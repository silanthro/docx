import json
import os
from pathlib import Path
from typing import List, Literal, TypedDict, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor


class NotAuthorizedError(Exception):
    pass


def _get_allowed_dir() -> list[Path]:
    allowed_dir_str = os.environ.get("ALLOWED_DIR")
    if not allowed_dir_str:
        return None
    if "[" in allowed_dir_str:
        return [Path(p).resolve() for p in json.loads(allowed_dir_str)]
    else:
        return [Path(allowed_dir_str).resolve()]


ALLOWED_DIR = _get_allowed_dir()


def _path_is_allowed(path: str):
    if ALLOWED_DIR is None:
        return True
    if any(Path(path).resolve() == dir for dir in ALLOWED_DIR):
        return True
    if not any(dir in Path(path).resolve().parents for dir in ALLOWED_DIR):
        raise NotAuthorizedError()


class RunData(TypedDict, total=False):
    text: str
    bold: bool
    italic: bool
    underline: Union[bool, str]
    strikethrough: bool
    double_strike: bool
    superscript: bool
    subscript: bool
    font_name: str
    font_size: float  # in pt
    font_color: str  # e.g., "FF0000"
    highlight: Literal[
        "YELLOW",
        "GREEN",
        "CYAN",
        "PINK",
        "BLUE",
        "RED",
        "GRAY",
        "DARK_YELLOW",
        "DARK_GREEN",
        "DARK_CYAN",
        "DARK_PINK",
        "DARK_BLUE",
        "DARK_RED",
        "DARK_GRAY",
        "BLACK",
        "WHITE",
        "TURQUOISE",
        "NONE",
        "AUTO",
    ]
    all_caps: bool
    small_caps: bool
    hidden: bool


class ParagraphData(TypedDict, total=False):
    style: str
    alignment: Literal["LEFT", "CENTER", "RIGHT", "JUSTIFY"]
    left_indent: float  # in pt
    right_indent: float
    first_line_indent: float
    line_spacing: float
    runs: List[RunData]


DocData = List[ParagraphData]


def read_docx(path: str) -> DocData:
    """
    Read a .docx file and extract content with formatting details.
    Excludes any style or formatting fields that are default (i.e., None or Normal).
    Returns a list of paragraphs, headings or tables, each including minimal style info.

    Args:
    - path (str): Absolute path to docx file

    Returns:
        A list of dictionaries, where each dictionary represents either a paragraph, heading, or a table

    """
    _path_is_allowed(path)
    doc = Document(path)
    content = []

    def extract_run_formatting(run) -> dict:
        font = run.font
        run_data = {"text": run.text}

        if font.bold:
            run_data["bold"] = True
        if font.italic:
            run_data["italic"] = True
        if font.underline:
            run_data["underline"] = font.underline
        if font.strike:
            run_data["strikethrough"] = True
        if font.double_strike:
            run_data["double_strike"] = True
        if font.superscript:
            run_data["superscript"] = True
        if font.subscript:
            run_data["subscript"] = True
        if font.name:
            run_data["font_name"] = font.name
        if font.size:
            run_data["font_size"] = font.size.pt
        if font.color and font.color.rgb:
            run_data["font_color"] = str(font.color.rgb)
        if font.highlight_color:
            run_data["highlight"] = font.highlight_color.name
        if font.all_caps:
            run_data["all_caps"] = True
        if font.small_caps:
            run_data["small_caps"] = True
        if font.hidden:
            run_data["hidden"] = True

        return run_data

    def extract_paragraph_formatting(para) -> dict:
        para_data = {}
        if para.style.name != "Normal":
            para_data["style"] = para.style.name
        if para.alignment is not None:
            para_data["alignment"] = para.alignment.name

        fmt = para.paragraph_format
        if fmt.left_indent:
            para_data["left_indent"] = fmt.left_indent.pt
        if fmt.right_indent:
            para_data["right_indent"] = fmt.right_indent.pt
        if fmt.first_line_indent:
            para_data["first_line_indent"] = fmt.first_line_indent.pt
        if fmt.line_spacing:
            para_data["line_spacing"] = fmt.line_spacing

        runs_data = [extract_run_formatting(run) for run in para.runs]
        para_data["runs"] = runs_data
        return para_data

    # Process document content
    for element in doc.element.body:
        if element.tag.endswith("tbl"):  # Table
            table_obj = None
            # Map the XML node back to its python-docx Table, so we can use table_obj.rows, table_obj.cell(), etc.
            for t in doc.tables:
                if t._element is element:
                    table_obj = t
                    break

            if table_obj:
                table = {"type": "table", "rows": []}

                # Process table cells
                vertical_merges = {}  # Track vertical merges

                for row_idx, row in enumerate(table_obj.rows):
                    table_row = []
                    col_offset = 0

                    for col_idx, cell in enumerate(row.cells):
                        # Get cell spans
                        tc = cell._tc
                        grid_span = tc.tcPr.xpath("./w:gridSpan")  # Horizontal merge
                        v_merge = tc.tcPr.xpath("./w:vMerge")  # Vertical merge

                        # Calculate actual column index accounting for previous spans
                        actual_col = col_idx + col_offset

                        # Extract cell content
                        cell_data = {"paragraphs": []}
                        for para in cell.paragraphs:
                            cell_data["paragraphs"].append(
                                extract_paragraph_formatting(para)
                            )

                        # Handle horizontal merge (gridSpan)
                        h_span = 1
                        if grid_span:
                            h_span = int(grid_span[0].val)
                            col_offset += h_span - 1

                        # Handle vertical merge (vMerge)
                        if v_merge:
                            merge_val = v_merge[0].val
                            if merge_val == "restart":
                                # Start of vertical merge
                                vertical_merges[(row_idx, actual_col)] = cell_data
                            elif (
                                not merge_val
                                and (row_idx - 1, actual_col) in vertical_merges
                            ):
                                # Reuse the cell data from the previous row
                                cell_data = vertical_merges[(row_idx - 1, actual_col)]
                                vertical_merges[(row_idx, actual_col)] = cell_data

                        # Add the cell data h_span times
                        for _ in range(h_span):
                            table_row.append(cell_data)

                    table["rows"].append(table_row)

                content.append(table)

        elif element.tag.endswith("p"):  # Paragraph, heading, or title
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para:
                para_data = extract_paragraph_formatting(para)
                para_data["type"] = "paragraph"
                content.append(para_data)

    return content


def write_docx(data: DocData, output_path: str, overwrite: bool = False):
    """
    Write a .docx file from structured data returned by `read_docx()`.

    Args:
    - data (list): List of dictionaries representing paragraphs or headings
    - output_path (str): Absolute path to save docx
    - overwrite (bool): Whether to overwrite if output_path exists, defaults to False

    Returns:
        "Saved file"

    """
    _path_is_allowed(output_path)

    if not overwrite and Path(output_path).exists():
        raise ValueError(
            f"File already exists at {output_path}. Set overwrite=True to overwrite"
        )

    doc = Document()

    for para_info in data:
        para = doc.add_paragraph()

        # Set paragraph style if given
        if "style" in para_info:
            para.style = para_info["style"]

        # Set alignment
        if "alignment" in para_info:
            para.alignment = WD_ALIGN_PARAGRAPH[para_info["alignment"]]

        # Set paragraph formatting
        fmt = para.paragraph_format
        if "left_indent" in para_info:
            fmt.left_indent = Pt(para_info["left_indent"])
        if "right_indent" in para_info:
            fmt.right_indent = Pt(para_info["right_indent"])
        if "first_line_indent" in para_info:
            fmt.first_line_indent = Pt(para_info["first_line_indent"])
        if "line_spacing" in para_info:
            fmt.line_spacing = para_info["line_spacing"]

        # Add runs
        for run_info in para_info["runs"]:
            run = para.add_run(run_info["text"])

            font = run.font
            if run_info.get("bold"):
                font.bold = True
            if run_info.get("italic"):
                font.italic = True
            if run_info.get("underline") is not None:
                font.underline = run_info["underline"]
            if run_info.get("strikethrough"):
                font.strike = True
            if run_info.get("double_strike"):
                font.double_strike = True
            if run_info.get("superscript"):
                font.superscript = True
            if run_info.get("subscript"):
                font.subscript = True
            if run_info.get("font_name"):
                font.name = run_info["font_name"]
            if run_info.get("font_size"):
                font.size = Pt(run_info["font_size"])
            if run_info.get("font_color"):
                font.color.rgb = RGBColor.from_string(run_info["font_color"])
            if run_info.get("highlight"):
                font.highlight_color = WD_COLOR_INDEX[run_info["highlight"]]
            if run_info.get("all_caps"):
                font.all_caps = True
            if run_info.get("small_caps"):
                font.small_caps = True
            if run_info.get("hidden"):
                font.hidden = True

    doc.save(output_path)
    return "Saved file"
